"""
Generates Servallab_QC_Technical_Briefing.docx
Run: python3 generate_briefing.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Styles helpers ────────────────────────────────────────────────────────────

def set_font(run, size=11, bold=False, color=None, mono=False):
    run.bold = bold
    run.font.size = Pt(size)
    if mono:
        run.font.name = "Courier New"
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1, color=(31, 73, 125)):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(*color)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        set_font(r, bold=True)
    p.add_run(text)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    set_font(run, size=9, mono=True)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F2F2F2")
    p._element.get_or_add_pPr().append(shading)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════════════════

title = doc.add_heading("Servallab — QC Engine Technical Briefing", 0)
for run in title.runs:
    run.font.color.rgb = RGBColor(31, 73, 125)

doc.add_paragraph("Prepared for: Pitch & Technical Q&A  |  June 2026")
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 1. WHAT THE SYSTEM IS
# ══════════════════════════════════════════════════════════════════════════════

heading("1. What the System Is")
body(
    "Servallab is a CATI (Computer-Assisted Telephone Interviewing) survey data "
    "Quality Control engine. You upload a raw survey dataset (CSV, Excel, or SPSS .sav), "
    "configure your checks, and the system automatically identifies every row of data that "
    "is suspicious, inconsistent, or potentially fabricated. Results come back as a "
    "structured report with flagged rows, severity levels, and per-interviewer summaries."
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. PIPELINE
# ══════════════════════════════════════════════════════════════════════════════

heading("2. The Pipeline")
body("Data flows through five stages:")
for step in [
    ("Upload", "CSV, Excel (.xlsx), or SPSS (.sav) file is received"),
    ("Load", "core/loader.py ingests the file into a structured table (DataFrame)"),
    ("Clean", "core/cleaner.py normalises nulls, coerces types (e.g. '18' → number 18)"),
    ("Rule Engine", "core/rule_engine.py reads config, assembles and runs all active checks"),
    ("Report", "Results sent to the React frontend as JSON; each tab shows one category of issues"),
]:
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(step[0] + " — ")
    r.bold = True
    p.add_run(step[1])

doc.add_paragraph()
body(
    "Every check follows the same contract: take a table in, return a list of "
    "bad rows with an explanation column (prefixed _) for why it was flagged."
)

# ══════════════════════════════════════════════════════════════════════════════
# 3. CHECKS
# ══════════════════════════════════════════════════════════════════════════════

heading("3. Every Check — Plain Language")

checks = [
    (
        "Missing Value Check",
        "Finds rows where required fields are blank. With a threshold (e.g. 0.2), only "
        "columns where more than 20% of values are missing get flagged. Without a threshold, "
        "any blank in a required column is flagged.",
        'JSON key: "missing_threshold"\nExample: { "missing_threshold": 0.2 }',
        "warning",
    ),
    (
        "High Missing Column Check",
        "Flags entire columns (not individual rows) where the missing rate exceeds the threshold. "
        "Automatically paired with Missing Value Check.",
        None,
        "info",
    ),
    (
        "Range Check",
        "For numeric columns, flags any value outside a configured [min, max] window.",
        '{ "range_rules": [\n  { "column": "age", "min": 18, "max": 99 }\n] }',
        "warning",
    ),
    (
        "Duration Check",
        "Flags interviews that are too short or too long. A 45-minute survey completed in "
        "90 seconds is almost certainly fabricated.",
        '{ "interview_duration": {\n  "enabled": true,\n  "column": "duration_minutes",\n'
        '  "min_expected": 5,\n  "max_expected": 120\n} }',
        "warning",
    ),
    (
        "Logic Check",
        "Enforces IF/THEN rules across columns. Example: 'IF age < 18 THEN salary must be null.' "
        "Every row where ALL if-conditions are true is checked against the then-conditions. "
        "Violation = row is flagged. Supported operators: >, <, >=, <=, ==, !=, is_null, "
        "not_null, in_list, not_in_list, is_numeric, is_string.",
        '{ "logic_rules": [{\n  "description": "Under-18 no salary",\n'
        '  "if_conditions":   [{ "column": "age",    "operator": "<",       "value": 18 }],\n'
        '  "then_conditions": [{ "column": "salary", "operator": "is_null" }]\n}] }',
        "critical",
    ),
    (
        "Duplicate Check",
        "Finds exact duplicate records based on specified columns (or all columns). "
        "Two rows with the same respondent ID and same answers — one is a re-entry.",
        '{ "duplicate_check": {\n  "enabled": true,\n  "subset_columns": ["respondent_id"]\n} }',
        "critical",
    ),
    (
        "Pattern Check (Regex)",
        "Uses a regular expression to validate the format of values. "
        "Example: phone numbers must match ^[0-9+\\-]+$ — any phone with letters gets flagged. "
        "A regex is a pattern-matching language: ^[0-9+\\-]+$ means 'only digits, plus signs, "
        "or hyphens, nothing else.'",
        '{ "pattern_rules": [{\n  "column": "phone",\n'
        '  "pattern": "^[0-9+\\\\-]+$",\n'
        '  "description": "Valid phone format"\n}] }',
        "warning",
    ),
    (
        "Anomaly Check (IQR)",
        "Finds statistical outliers in numeric columns using the Interquartile Range method. "
        "IQR = Q3 - Q1 (the spread of the middle 50%). "
        "Lower bound = Q1 - (multiplier x IQR). Upper bound = Q3 + (multiplier x IQR). "
        "Any value outside [lower, upper] is an outlier. "
        "Multiplier 1.5 = standard (Tukey fence); 3.0 = extreme outliers only. "
        "IQR is preferred over mean±SD because it is robust to the very outliers being detected.",
        '{ "anomaly_check": {\n  "enabled": true,\n  "columns": ["age", "income"],\n  "multiplier": 1.5\n} }',
        "info",
    ),
    (
        "Straightlining Check",
        "Detects respondents who gave the same answer to every question in a grid — a sign "
        "of not engaging. Threshold = proportion of identical answers needed to flag. "
        "0.9 = 90% same answer. Outputs _sl_score (exact proportion) and _sl_modal_answer "
        "(which answer they kept giving). Can also summarise per interviewer.",
        '{ "straightlining": {\n  "enabled": true,\n  "question_columns": ["Q1","Q2","Q3"],\n'
        '  "threshold": 0.9,\n  "interviewer_column": "int_id",\n  "min_questions": 3\n} }',
        "warning",
    ),
    (
        "Interviewer Duration Check",
        "Groups interviews by interviewer, computes each interviewer's average duration, "
        "then uses IQR to find interviewers whose average is an outlier. "
        "Flags as 'too_fast' or 'too_slow'. Only includes interviewers with >= min_interviews.",
        '{ "interviewer_duration_check": {\n  "enabled": true,\n'
        '  "interviewer_column": "int_id",\n  "duration_column": "duration_minutes",\n'
        '  "multiplier": 1.5,\n  "min_interviews": 3\n} }',
        "warning",
    ),
    (
        "Interviewer Productivity Check",
        "Flags interviewers with significantly more or fewer interviews than peers (IQR on counts). "
        "unusually_high = possible fabrication; unusually_low = low effort or data entry issues.",
        '{ "interviewer_productivity_check": {\n  "enabled": true,\n'
        '  "interviewer_column": "int_id",\n  "multiplier": 1.5\n} }',
        "warning",
    ),
    (
        "Consent / Eligibility Check",
        "Checks that respondents disqualified by a screener question have no data in subsequent "
        "questions. If consent != 'Yes' but Q2–Q20 are filled in, the skip logic failed or "
        "data was fabricated.",
        '{ "consent_eligibility_check": {\n  "enabled": true,\n'
        '  "screener_column": "consent",\n  "disqualify_operator": "!=",\n'
        '  "disqualify_value": "Yes",\n  "subsequent_columns": ["Q2","Q3","Q4"]\n} }',
        "critical",
    ),
    (
        "Fabrication Check",
        "Three sub-checks: (A) Sequential IDs — runs of 5+ consecutive respondent IDs in a "
        "random sample are suspicious. (B) Low variance — an interviewer's standard deviation "
        "on a numeric question is less than 10% of the global std dev, suggesting all their "
        "respondents gave near-identical numbers. (C) Repeated numeric blocks.",
        '{ "fabrication_check": {\n  "enabled": true,\n  "id_column": "respondent_id",\n'
        '  "numeric_columns": ["Q10","Q11"],\n  "interviewer_column": "int_id",\n'
        '  "variance_threshold": 0.1,\n  "sequence_run_length": 5\n} }',
        "critical",
    ),
    (
        "Near Duplicate Check",
        "Finds near-duplicates exact checks miss. (A) Shared unique identifiers — same phone "
        "or email under multiple respondent IDs (interviewer reused a contact). "
        "(B) Repeated demographic combo — same age+gender+location appearing more than "
        "max_combo_count times.",
        '{ "near_duplicate_check": {\n  "enabled": true,\n  "id_column": "respondent_id",\n'
        '  "unique_columns": ["phone","email"],\n  "combo_columns": ["age","gender","location"],\n'
        '  "max_combo_count": 3\n} }',
        "warning",
    ),
    (
        "Verbatim Quality Check (AI)",
        "Sends open-ended responses to Groq AI (LLaMA model) for scoring. "
        "Scores grammar, coherence, relevance, length_quality (1–5 each). "
        "Flags gibberish, copy-paste, too_short responses. "
        "Batches 10 responses per API call for efficiency. "
        "Uses server key (GROQ_API_KEY env var) with fallback to user's personal key.",
        '{ "verbatim_check": {\n  "enabled": true,\n  "verbatim_columns": ["Q_open"],\n'
        '  "model": "llama-3.1-8b-instant",\n  "min_score": 2,\n  "sample_size": 100\n} }',
        "warning",
    ),
]

for name, desc, json_ex, severity in checks:
    heading(name, level=2)
    body(desc)
    if json_ex:
        doc.add_paragraph()
        body("JSON config:")
        code_block(json_ex)
    sev_colors = {"critical": (192, 0, 0), "warning": (198, 109, 0), "info": (31, 73, 125)}
    p = doc.add_paragraph("Severity: ")
    r = p.add_run(severity.upper())
    r.bold = True
    r.font.color.rgb = RGBColor(*sev_colors[severity])
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 4. SEVERITY LEVELS
# ══════════════════════════════════════════════════════════════════════════════

heading("4. Severity Levels")
add_table(
    ["Level", "Meaning", "Examples"],
    [
        ["CRITICAL", "Data cannot be trusted — structural violation",
         "Logic violations, fabrication, consent breaches, exact duplicates"],
        ["WARNING", "Suspicious — needs human review",
         "Straightlining, duration anomalies, verbatim quality, near duplicates"],
        ["INFO", "Data quality signal — not necessarily wrong",
         "High missing column rate, statistical anomalies"],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# 5. THRESHOLDS — WHAT THEY ARE
# ══════════════════════════════════════════════════════════════════════════════

heading("5. What a Threshold Is")
body(
    "A threshold is a configurable cut-off that decides when something gets flagged. "
    "Every check has at least one. All thresholds are configurable — a short mobile survey "
    "has different valid duration bounds than a 60-minute omnibus. The engine adapts to the "
    "study design, not the other way around."
)
doc.add_paragraph()
add_table(
    ["Check", "Threshold parameter", "Meaning"],
    [
        ["Missing Value", "threshold: 0.2", "Flag if >20% blank"],
        ["Anomaly / IQR checks", "multiplier: 1.5", "How far outside middle 50% = outlier"],
        ["Straightlining", "threshold: 0.9", "Flag if 90%+ of answers are the same"],
        ["Fabrication", "variance_threshold: 0.1", "Flag if std dev < 10% of global"],
        ["Near Duplicate", "max_combo_count: 3", "Flag demographic combo if it appears >3×"],
        ["Duration", "min_minutes / max_minutes", "Hard floor and ceiling for interview length"],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# 6. IQR EXPLAINED
# ══════════════════════════════════════════════════════════════════════════════

heading("6. IQR — The Outlier Detection Method")
body("IQR = Interquartile Range = the spread of the middle 50% of your data.")
doc.add_paragraph()

steps = [
    "Sort all values in the column",
    "Q1 = value at the 25th percentile (bottom quarter)",
    "Q3 = value at the 75th percentile (top quarter)",
    "IQR = Q3 − Q1",
    "Lower bound = Q1 − (multiplier × IQR)",
    "Upper bound = Q3 + (multiplier × IQR)",
    "Any value outside [lower, upper] is an outlier and gets flagged",
]
for s in steps:
    doc.add_paragraph(s, style="List Number")

doc.add_paragraph()
body(
    "Example: Interview durations — Q1=15 min, Q3=30 min, IQR=15. "
    "With multiplier 1.5: bounds = [15−22.5, 30+22.5] = [−7.5, 52.5]. "
    "Any interview over 52.5 minutes is flagged. (Negative lower bound means no short outliers "
    "since durations cannot be negative.)"
)
doc.add_paragraph()
body(
    "Why IQR instead of mean ± standard deviation? "
    "IQR is robust to extremes — a few fabricated values at the top won't distort it. "
    "Standard deviation is inflated by the very values being detected, which widens the fence "
    "and misses the outliers. IQR is the industry standard (Tukey, 1977)."
)

# ══════════════════════════════════════════════════════════════════════════════
# 7. PATTERN / REGEX CHECK — JSON FORMAT
# ══════════════════════════════════════════════════════════════════════════════

heading("7. Pattern (Regex) Check — Full JSON Reference")
body(
    "The pattern check is triggered by the key 'pattern_rules' in the config JSON. "
    "Each rule targets one column and provides a regex pattern the values must match."
)
doc.add_paragraph()
code_block(
    '{\n'
    '  "pattern_rules": [\n'
    '    {\n'
    '      "column": "phone",\n'
    '      "pattern": "^[0-9+\\\\-]+$",\n'
    '      "description": "Valid phone — digits, +, - only"\n'
    '    },\n'
    '    {\n'
    '      "column": "email",\n'
    '      "pattern": "^[\\\\w.+-]+@[\\\\w-]+\\\\.[\\\\w.]+$",\n'
    '      "description": "Valid email format"\n'
    '    },\n'
    '    {\n'
    '      "column": "id_number",\n'
    '      "pattern": "^[A-Z]{2}[0-9]{6}$",\n'
    '      "description": "National ID — 2 uppercase letters then 6 digits"\n'
    '    }\n'
    '  ]\n'
    '}'
)
doc.add_paragraph()
body("Common ready-to-use regex patterns:")
add_table(
    ["What to validate", "Pattern", "Meaning"],
    [
        ["Digits only", "^[0-9]+$", "Only numbers, nothing else"],
        ["Phone (digits/+/-)", "^[0-9+\\-]+$", "Phone numbers in any international format"],
        ["Email", "^[\\w.+-]+@[\\w-]+\\.[\\w.]+$", "Standard email format"],
        ["Uppercase letters only", "^[A-Z]+$", "All caps, no digits or spaces"],
        ["Date YYYY-MM-DD", "^\\d{4}-\\d{2}-\\d{2}$", "ISO date format"],
        ["SA ID Number (13 digits)", "^\\d{13}$", "Exactly 13 digits"],
        ["2-letter country code", "^[A-Z]{2}$", "e.g. ZA, US, GB"],
    ]
)
body(
    "In the UI (Config tab), pattern rules can also be added through the 'Pattern Rules' "
    "form — no JSON editing required. Enter the column name, paste the regex, "
    "add a description, and click Add."
)

# ══════════════════════════════════════════════════════════════════════════════
# 8. TECHNICAL Q&A
# ══════════════════════════════════════════════════════════════════════════════

heading("8. Likely Technical Questions & Answers")

qas = [
    (
        "Why IQR and not mean ± 2 standard deviations?",
        "The IQR method is robust — outliers don't distort it. Standard deviation is inflated "
        "by the very values being detected, making the fence too wide and missing real outliers. "
        "IQR is the industry standard (Tukey, 1977)."
    ),
    (
        "Can the system handle SPSS .sav files?",
        "Yes. The loader uses pyreadstat to ingest SAV files natively, preserving value labels."
    ),
    (
        "What happens to uploaded data?",
        "Files are stored in a temporary directory and auto-deleted after 1 hour. "
        "Nothing is persisted beyond the session."
    ),
    (
        "Can I define logic rules without coding?",
        "Yes. Rules are JSON objects editable in the Config tab. A non-technical supervisor "
        "can write IF/THEN conditions without touching Python."
    ),
    (
        "How does the AI verbatim check work without sending all data to OpenAI?",
        "It uses Groq — a different inference provider running open-source LLaMA models. "
        "Only the text of open-ended answers is sent (not the full dataset), batched 10 per call."
    ),
    (
        "How fast is it on a large dataset?",
        "The straightlining check uses vectorised NumPy operations — 3–5× faster than row-by-row "
        "Python loops. IQR checks run at Pandas aggregate speed. A 10,000-row dataset with "
        "all checks enabled runs in seconds."
    ),
    (
        "What if an interviewer only has 1 or 2 interviews?",
        "The duration and productivity checks have a min_interviews parameter (default: 3). "
        "Interviewers below this threshold are excluded from IQR analysis to avoid false positives."
    ),
    (
        "Can multiple pattern rules be applied to the same column?",
        "Yes — just add multiple objects to the pattern_rules array, each targeting the same column. "
        "Each rule is evaluated independently and violations are reported separately."
    ),
]

for q, a in qas:
    p = doc.add_paragraph()
    r = p.add_run("Q: " + q)
    r.bold = True
    doc.add_paragraph("A: " + a)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════

out = "/home/dvock/QC-Engine-desktop/Servallab_QC_Technical_Briefing.docx"
doc.save(out)
print(f"Saved: {out}")
